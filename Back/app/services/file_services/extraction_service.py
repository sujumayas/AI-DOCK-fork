"""
Text Extraction Service for AI Dock

Atomic service responsible for extracting text content from various file formats:
- PDF text extraction using PyPDF2
- DOCX text extraction using python-docx and docx2txt
- DOC handling with informative messages
- Plain text processing
- Content cleaning and normalization

🎓 LEARNING: Strategy Pattern Implementation
==========================================
This service implements the Strategy pattern for different file types:
- Each file type has its own extraction strategy
- Common interface for all extraction methods
- Easy to add new file type support
- Follows integration guide's modular patterns
"""

import re
from io import BytesIO
from typing import Tuple, Optional, Dict, Callable

# Text extraction libraries (optional imports)
try:
    import PyPDF2
    PDF_EXTRACTION_AVAILABLE = True
except ImportError:
    PDF_EXTRACTION_AVAILABLE = False

try:
    import docx2txt
    from docx import Document as DocxDocument
    DOCX_EXTRACTION_AVAILABLE = True
except ImportError:
    DOCX_EXTRACTION_AVAILABLE = False

# Internal imports
from ...schemas.file_upload import AllowedFileType


class TextExtractionService:
    """
    Atomic service for text extraction from various file formats.
    
    Following integration guide patterns:
    - Strategy pattern for different file types
    - Comprehensive error handling
    - Clean, normalized output
    - No side effects (pure extraction)
    """
    
    def __init__(self):
        """Initialize extraction service with format-specific strategies."""
        # Register extraction strategies for each file type
        self._extraction_strategies: Dict[str, Callable] = {
            AllowedFileType.PDF.value: self._extract_from_pdf,
            AllowedFileType.DOCX.value: self._extract_from_docx,
            AllowedFileType.DOC.value: self._extract_from_doc,
            'text/plain': self._extract_from_text,
            'text/markdown': self._extract_from_text,
            'text/csv': self._extract_from_text,
            'application/json': self._extract_from_text,
            'text/x-python': self._extract_from_text,
            'text/javascript': self._extract_from_text,
            'text/html': self._extract_from_text,
            'text/css': self._extract_from_text,
            'application/xml': self._extract_from_text,
            'text/xml': self._extract_from_text,
        }
    
    # =============================================================================
    # MAIN EXTRACTION ENTRY POINT
    # =============================================================================
    
    def extract_text_content(self, content_bytes: bytes, filename: str, content_type: str) -> Tuple[str, Optional[str]]:
        """
        Extract text content from file bytes using appropriate strategy.
        
        🎓 LEARNING: Strategy Pattern
        ============================
        Instead of a large if/elif chain, we use a strategy dictionary:
        - Cleaner code structure
        - Easy to extend with new formats
        - Testable individual strategies
        - Follows open/closed principle
        
        Args:
            content_bytes: File content as bytes
            filename: Original filename for error messages
            content_type: MIME type to determine extraction strategy
            
        Returns:
            Tuple of (extracted_text, error_message)
        """
        try:
            # Get appropriate extraction strategy
            extraction_strategy = self._extraction_strategies.get(content_type)
            
            if not extraction_strategy:
                return "", f"Text extraction not supported for file type: {content_type}"
            
            # Execute extraction strategy
            extracted_text, error_message = extraction_strategy(content_bytes, filename)
            
            if error_message:
                # Log warning but don't fail upload
                print(f"Text extraction warning for {filename}: {error_message}")
                return "", None  # Return empty text, no error
            
            # Clean and normalize extracted text
            if extracted_text:
                cleaned_text = self._clean_extracted_text(extracted_text)
                return cleaned_text, None
            else:
                return "", None
                
        except Exception as e:
            error_msg = f"Text extraction failed for {filename}: {str(e)}"
            print(f"Text extraction error: {error_msg}")
            return "", None  # Return empty text, don't fail upload
    
    # =============================================================================
    # PDF EXTRACTION STRATEGY
    # =============================================================================
    
    def _extract_from_pdf(self, content_bytes: bytes, filename: str) -> Tuple[str, Optional[str]]:
        """
        Extract text from PDF file content using PyPDF2.
        
        🎓 LEARNING: PDF Text Extraction
        ===============================
        PDFs can contain:
        - Selectable text (easy to extract)
        - Scanned images (need OCR - future enhancement)
        - Password protection (not supported)
        - Corrupted content (handle gracefully)
        
        We use PyPDF2 for basic text extraction. For advanced features
        like OCR, we'd need additional libraries like Tesseract.
        
        Args:
            content_bytes: PDF file content as bytes
            filename: Original filename for error messages
            
        Returns:
            Tuple of (extracted_text, error_message)
        """
        if not PDF_EXTRACTION_AVAILABLE:
            return "", "PDF text extraction not available (PyPDF2 not installed)"
        
        try:
            pdf_file = BytesIO(content_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Check if PDF is encrypted
            if pdf_reader.is_encrypted:
                return "", f"PDF {filename} is password protected and cannot be processed"
            
            # Extract text from all pages
            extracted_text = ""
            total_pages = len(pdf_reader.pages)
            
            if total_pages == 0:
                return "", f"PDF {filename} contains no pages"
            
            # Process each page
            successful_pages = 0
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():  # Only add non-empty text
                        extracted_text += f"\\n--- Page {page_num + 1} ---\\n"
                        extracted_text += page_text.strip() + "\\n"
                        successful_pages += 1
                except Exception as page_error:
                    # Continue with other pages if one fails
                    print(f"Warning: Failed to extract text from page {page_num + 1} of {filename}: {page_error}")
                    continue
            
            if not extracted_text.strip():
                return "", f"PDF {filename} appears to contain no extractable text (may be scanned images)"
            
            # Add extraction summary
            if successful_pages < total_pages:
                extracted_text += f"\\n\\n[Note: Successfully extracted text from {successful_pages} of {total_pages} pages]"
            
            return extracted_text, None
            
        except Exception as e:
            error_msg = f"Failed to extract text from PDF {filename}: {str(e)}"
            return "", error_msg
    
    # =============================================================================
    # DOCX EXTRACTION STRATEGY
    # =============================================================================
    
    def _extract_from_docx(self, content_bytes: bytes, filename: str) -> Tuple[str, Optional[str]]:
        """
        Extract text from modern Word (.docx) file content.
        
        🎓 LEARNING: DOCX Text Extraction
        =================================
        .docx files are ZIP archives containing XML files.
        We can extract text using:
        - docx2txt: Simple, fast text extraction
        - python-docx: Full document object model with structure
        
        Strategy: Try docx2txt first (simpler), fallback to python-docx
        
        Args:
            content_bytes: DOCX file content as bytes
            filename: Original filename for error messages
            
        Returns:
            Tuple of (extracted_text, error_message)
        """
        if not DOCX_EXTRACTION_AVAILABLE:
            return "", "DOCX text extraction not available (python-docx or docx2txt not installed)"
        
        try:
            # Strategy 1: Try docx2txt first (simpler and more reliable)
            extracted_text = self._try_docx2txt_extraction(content_bytes)
            
            if extracted_text and extracted_text.strip():
                return extracted_text, None
            
            # Strategy 2: Fallback to python-docx (more comprehensive)
            extracted_text = self._try_python_docx_extraction(content_bytes)
            
            if extracted_text and extracted_text.strip():
                return extracted_text, None
            
            return "", f"Word document {filename} appears to contain no extractable text"
            
        except Exception as e:
            error_msg = f"Failed to extract text from Word document {filename}: {str(e)}"
            return "", error_msg
    
    def _try_docx2txt_extraction(self, content_bytes: bytes) -> str:
        """Try extracting text using docx2txt library."""
        try:
            docx_file = BytesIO(content_bytes)
            extracted_text = docx2txt.process(docx_file)
            return extracted_text if extracted_text else ""
        except Exception as e:
            print(f"docx2txt extraction failed: {e}")
            return ""
    
    def _try_python_docx_extraction(self, content_bytes: bytes) -> str:
        """Try extracting text using python-docx library."""
        try:
            docx_file = BytesIO(content_bytes)
            doc = DocxDocument(docx_file)
            
            extracted_text = ""
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    extracted_text += para.text + "\\n"
            
            # Extract text from tables
            table_content = self._extract_table_content(doc)
            if table_content:
                extracted_text += "\\n" + table_content
            
            return extracted_text
            
        except Exception as e:
            print(f"python-docx extraction failed: {e}")
            return ""
    
    def _extract_table_content(self, doc) -> str:
        """Extract structured content from document tables."""
        table_text = ""
        
        for table in doc.tables:
            table_text += "\\n--- Table ---\\n"
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    table_text += " | ".join(row_text) + "\\n"
        
        return table_text
    
    # =============================================================================
    # DOC EXTRACTION STRATEGY
    # =============================================================================
    
    def _extract_from_doc(self, content_bytes: bytes, filename: str) -> Tuple[str, Optional[str]]:
        """
        Handle legacy Word (.doc) file content.
        
        🎓 LEARNING: Legacy Format Handling
        ==================================
        .doc files use the OLE compound document format.
        They're much harder to parse than .docx files.
        
        For now, we provide helpful guidance to users.
        Future enhancements could include:
        - python-doc library integration
        - antiword command-line tool
        - LibreOffice API conversion
        
        Args:
            content_bytes: DOC file content as bytes
            filename: Original filename for error messages
            
        Returns:
            Tuple of (extracted_text, error_message)
        """
        # For now, we don't support .doc extraction
        # This is a complex format that requires specialized libraries
        error_msg = (
            f"Legacy Word document format (.doc) is not yet supported for text extraction. "
            f"Please convert {filename} to .docx format and upload again."
        )
        return "", error_msg
    
    # =============================================================================
    # TEXT FILE EXTRACTION STRATEGY
    # =============================================================================
    
    def _extract_from_text(self, content_bytes: bytes, filename: str) -> Tuple[str, Optional[str]]:
        """
        Extract text from plain text files with encoding detection.
        
        🎓 LEARNING: Text Encoding Handling
        ==================================
        Text files can have different encodings:
        - UTF-8 (most common, preferred)
        - UTF-8 with BOM
        - Latin-1/CP1252 (legacy)
        - Others (regional encodings)
        
        Strategy: Try common encodings in order of preference
        
        Args:
            content_bytes: Text file content as bytes
            filename: Original filename for error messages
            
        Returns:
            Tuple of (extracted_text, error_message)
        """
        try:
            # Try encodings in order of preference
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text_content = content_bytes.decode(encoding)
                    return text_content, None
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, decode with error replacement
            text_content = content_bytes.decode('utf-8', errors='replace')
            return text_content, f"Warning: {filename} contains non-UTF-8 characters, some may be replaced"
            
        except Exception as e:
            error_msg = f"Failed to decode text file {filename}: {str(e)}"
            return "", error_msg
    
    # =============================================================================
    # TEXT CLEANING AND NORMALIZATION
    # =============================================================================
    
    def _clean_extracted_text(self, text: str) -> str:
        """
        Clean up extracted text by removing excessive whitespace and normalizing.
        
        🎓 LEARNING: Text Normalization
        ==============================
        Extracted text often contains:
        - Multiple consecutive spaces
        - Excessive line breaks
        - Special characters
        - Formatting artifacts
        
        We clean it up for better readability and LLM processing.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned and normalized text
        """
        if not text:
            return ""
        
        # Step 1: Normalize line endings
        text = text.replace('\\r\\n', '\\n').replace('\\r', '\\n')
        
        # Step 2: Process lines individually
        lines = []
        for line in text.split('\\n'):
            # Remove extra spaces within lines
            cleaned_line = ' '.join(line.split())
            if cleaned_line:  # Only keep non-empty lines
                lines.append(cleaned_line)
        
        # Step 3: Join lines with single newlines
        cleaned_text = '\\n'.join(lines)
        
        # Step 4: Remove excessive consecutive newlines (max 2)
        cleaned_text = re.sub(r'\\n{3,}', '\\n\\n', cleaned_text)
        
        # Step 5: Remove leading/trailing whitespace
        cleaned_text = cleaned_text.strip()
        
        # Step 6: Limit content size to prevent token overflow
        max_chars = 50000  # ~50k characters for LLM context
        if len(cleaned_text) > max_chars:
            cleaned_text = cleaned_text[:max_chars] + "\\n\\n[Content truncated - file is larger than 50k characters]"
        
        return cleaned_text
    
    # =============================================================================
    # UTILITY METHODS
    # =============================================================================
    
    def get_supported_types(self) -> list:
        """Get list of file types that support text extraction."""
        return list(self._extraction_strategies.keys())
    
    def is_extraction_available_for_type(self, content_type: str) -> bool:
        """Check if text extraction is available for given content type."""
        if content_type not in self._extraction_strategies:
            return False
        
        # Check if required libraries are available
        if content_type == AllowedFileType.PDF.value:
            return PDF_EXTRACTION_AVAILABLE
        elif content_type in [AllowedFileType.DOCX.value, AllowedFileType.DOC.value]:
            return DOCX_EXTRACTION_AVAILABLE
        else:
            # Text files don't need special libraries
            return True
    
    def get_extraction_requirements(self) -> Dict[str, Dict[str, str]]:
        """Get information about extraction requirements for different file types."""
        return {
            AllowedFileType.PDF.value: {
                "library": "PyPDF2",
                "available": str(PDF_EXTRACTION_AVAILABLE),
                "features": "Basic text extraction, password detection"
            },
            AllowedFileType.DOCX.value: {
                "library": "python-docx, docx2txt",
                "available": str(DOCX_EXTRACTION_AVAILABLE),
                "features": "Full text and table extraction"
            },
            AllowedFileType.DOC.value: {
                "library": "Not implemented",
                "available": "False",
                "features": "Legacy format - conversion recommended"
            },
            "text/*": {
                "library": "Built-in",
                "available": "True",
                "features": "Multi-encoding support"
            }
        }
    
    def estimate_extraction_time(self, file_size_bytes: int, content_type: str) -> str:
        """Estimate extraction time based on file size and type."""
        size_mb = file_size_bytes / (1024 * 1024)
        
        if content_type == AllowedFileType.PDF.value:
            # PDFs take longer due to structure parsing
            if size_mb < 1:
                return "< 5 seconds"
            elif size_mb < 5:
                return "5-15 seconds"
            else:
                return "15-30 seconds"
        elif content_type == AllowedFileType.DOCX.value:
            # DOCX processing is moderately fast
            if size_mb < 2:
                return "< 3 seconds"
            elif size_mb < 10:
                return "3-10 seconds"
            else:
                return "10-20 seconds"
        else:
            # Text files are very fast
            return "< 2 seconds"
