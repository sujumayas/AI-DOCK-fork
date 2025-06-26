"""
Word document processor implementation.

This module handles Microsoft Word documents (.docx, .doc) with robust text extraction,
structure preservation, and error handling for various Word-specific issues.

🎓 LEARNING: Word Document Processing
===================================
Word document processing challenges:
- Different formats (.doc vs .docx)
- Password-protected documents
- Complex formatting and embedded objects
- Large documents with lots of content
- Macros and embedded scripts
"""

import logging
from typing import Dict, Any, Optional
from pathlib import Path
import io

try:
    from docx import Document
    from docx.oxml.exceptions import InvalidXmlError
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

try:
    import zipfile
    HAS_ZIPFILE = True
except ImportError:
    HAS_ZIPFILE = False

from .base_processor import BaseFileProcessor
from ..types import ContentType, WordMetadata
from ..config import config
from ..exceptions import (
    WordProcessingError,
    WordPasswordProtectedError,
    WordCorruptedError,
    WordUnsupportedFormatError,
    WordMacroContentError,
    UnsupportedFileTypeError
)
from ...models.file_upload import FileUpload

logger = logging.getLogger(__name__)


class WordProcessor(BaseFileProcessor):
    """
    Microsoft Word document processor with comprehensive text extraction.
    
    🎓 LEARNING: Word Document Structure
    ===================================
    Word documents are complex:
    - .docx files are ZIP archives containing XML
    - .doc files use older binary format (not supported here)
    - Documents can contain embedded objects, images, tables
    - Text formatting affects extraction quality
    """
    
    SUPPORTED_MIME_TYPES = {
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',  # .docx
        'application/msword'  # .doc (limited support)
    }
    
    def __init__(self):
        super().__init__()
        self.check_dependencies()
    
    def check_dependencies(self) -> None:
        """Check if required Word processing libraries are available."""
        if not HAS_PYTHON_DOCX:
            raise UnsupportedFileTypeError(
                "Word document processing requires python-docx library. "
                "Install with: pip install python-docx"
            )
    
    def can_process(self, mime_type: str) -> bool:
        """Check if this processor can handle the given MIME type."""
        return mime_type in self.SUPPORTED_MIME_TYPES
    
    def get_content_type(self, mime_type: str) -> ContentType:
        """Get the appropriate ContentType for Word documents."""
        return ContentType.STRUCTURED_DATA
    
    def validate_format_specific(self, file_upload: FileUpload) -> None:
        """Word document-specific validation."""
        # Check if file size is within Word limits
        if file_upload.file_size > self.config.word.max_word_size:
            raise WordProcessingError(
                f"Word document too large: {file_upload.file_size} bytes "
                f"(max {self.config.word.max_word_size})",
                file_upload.id
            )
        
        # Check if it's .doc format (limited support)
        if file_upload.mime_type == 'application/msword':
            raise WordUnsupportedFormatError(
                f"Legacy .doc format not supported for {file_upload.filename}. "
                f"Please convert to .docx format.",
                file_upload.id
            )
    
    async def extract_content(self, file_upload: FileUpload) -> str:
        """
        Extract text content from Word document.
        
        🎓 LEARNING: Word Text Extraction Strategy
        =========================================
        Word extraction approach:
        1. Use python-docx to read .docx files
        2. Extract paragraphs, tables, headers/footers
        3. Preserve document structure (headings, lists)
        4. Handle embedded objects gracefully
        5. Extract metadata about document structure
        
        Args:
            file_upload: FileUpload model instance
            
        Returns:
            Extracted and formatted text content
            
        Raises:
            WordProcessingError: If extraction fails
        """
        try:
            # Get file content as bytes
            word_bytes = await self._get_word_bytes(file_upload)
            
            # Check if it's a valid .docx file
            await self._validate_docx_format(word_bytes, file_upload)
            
            # Extract content using python-docx
            return await self._extract_with_python_docx(word_bytes, file_upload)
            
        except WordProcessingError:
            # Re-raise our custom errors
            raise
        except Exception as e:
            self.logger.error(f"Unexpected error extracting Word doc {file_upload.filename}: {e}")
            raise WordProcessingError(
                f"Failed to extract content from Word document: {str(e)}",
                file_upload.id
            )
    
    async def _get_word_bytes(self, file_upload: FileUpload) -> bytes:
        """Get Word document content as bytes from file_upload."""
        if hasattr(file_upload, 'file_path') and file_upload.file_path:
            # File stored on disk
            file_path = Path(file_upload.file_path)
            if not file_path.exists():
                raise WordProcessingError(
                    f"Word document not found: {file_path}",
                    file_upload.id
                )
            return file_path.read_bytes()
        
        elif hasattr(file_upload, 'file_data') and file_upload.file_data:
            # File stored as binary data in database
            return file_upload.file_data
        
        else:
            raise WordProcessingError(
                f"No accessible content for Word document {file_upload.filename}",
                file_upload.id
            )
    
    async def _validate_docx_format(self, word_bytes: bytes, file_upload: FileUpload) -> None:
        """Validate that the file is a proper .docx format."""
        try:
            # .docx files are ZIP archives - check if it's a valid ZIP
            if HAS_ZIPFILE:
                with zipfile.ZipFile(io.BytesIO(word_bytes), 'r') as zip_file:
                    # Check for required .docx files
                    required_files = ['word/document.xml', '[Content_Types].xml']
                    file_list = zip_file.namelist()
                    
                    for required_file in required_files:
                        if required_file not in file_list:
                            raise WordCorruptedError(
                                f"Word document {file_upload.filename} is missing required component: {required_file}",
                                file_upload.id
                            )
        except zipfile.BadZipFile:
            raise WordCorruptedError(
                f"Word document {file_upload.filename} is not a valid .docx file (corrupted ZIP archive)",
                file_upload.id
            )
        except WordProcessingError:
            raise
        except Exception as e:
            self.logger.warning(f"Could not validate .docx format for {file_upload.filename}: {e}")
            # Continue anyway - python-docx might still be able to handle it
    
    async def _extract_with_python_docx(self, word_bytes: bytes, file_upload: FileUpload) -> str:
        """
        Extract text using python-docx library.
        
        🎓 LEARNING: python-docx Text Extraction
        =======================================
        python-docx provides structured access to:
        - Document paragraphs with formatting
        - Tables with cell content
        - Headers and footers
        - Document properties and metadata
        - Styles and structural elements
        """
        try:
            # Load document from bytes
            doc = Document(io.BytesIO(word_bytes))
            
            extracted_parts = []
            
            # Extract document properties first
            if doc.core_properties.title:
                extracted_parts.append(f"Title: {doc.core_properties.title}")
            
            if doc.core_properties.subject:
                extracted_parts.append(f"Subject: {doc.core_properties.subject}")
            
            if extracted_parts:
                extracted_parts.append("=" * 50)  # Separator
            
            # Extract paragraphs
            paragraph_count = 0
            current_text_length = 0
            
            for paragraph in doc.paragraphs:
                paragraph_text = paragraph.text.strip()
                
                if paragraph_text:  # Skip empty paragraphs
                    # Check for heading styles
                    if paragraph.style.name.startswith('Heading'):
                        # Format headings with appropriate markers
                        level = self._get_heading_level(paragraph.style.name)
                        heading_marker = "#" * level
                        extracted_parts.append(f"\n{heading_marker} {paragraph_text}\n")
                    else:
                        extracted_parts.append(paragraph_text)
                    
                    paragraph_count += 1
                    current_text_length += len(paragraph_text)
                    
                    # Check size limits
                    if current_text_length > self.config.word.max_word_text:
                        self.logger.info(
                            f"Stopping Word extraction at paragraph {paragraph_count} "
                            f"due to size limit ({current_text_length} chars)"
                        )
                        extracted_parts.append(
                            f"\n[Document continues... stopped at {paragraph_count} paragraphs "
                            f"due to size limit]"
                        )
                        break
            
            # Extract tables
            table_count = 0
            for table in doc.tables:
                if current_text_length > self.config.word.max_word_text:
                    break
                
                table_count += 1
                extracted_parts.append(f"\n--- Table {table_count} ---")
                
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_cells.append(cell_text)
                    
                    if row_cells:
                        extracted_parts.append(" | ".join(row_cells))
                        current_text_length += sum(len(cell) for cell in row_cells)
                        
                        if current_text_length > self.config.word.max_word_text:
                            extracted_parts.append("[Table truncated due to size limit]")
                            break
                
                extracted_parts.append("")  # Empty line after table
            
            if not extracted_parts or all(not part.strip() for part in extracted_parts):
                raise WordProcessingError(
                    f"Word document {file_upload.filename} contains no extractable text",
                    file_upload.id
                )
            
            # Join all parts
            extracted_text = '\n'.join(extracted_parts)
            
            self.logger.info(
                f"python-docx extracted {len(extracted_text)} chars from "
                f"{paragraph_count} paragraphs and {table_count} tables "
                f"in {file_upload.filename}"
            )
            
            return extracted_text
            
        except InvalidXmlError as e:
            raise WordCorruptedError(
                f"Word document {file_upload.filename} contains invalid XML: {str(e)}",
                file_upload.id
            )
        except PermissionError as e:
            raise WordPasswordProtectedError(
                f"Word document {file_upload.filename} appears to be password-protected",
                file_upload.id
            )
        except WordProcessingError:
            # Re-raise our custom errors
            raise
        except Exception as e:
            # Handle other potential errors
            error_str = str(e).lower()
            
            if "password" in error_str or "encrypted" in error_str:
                raise WordPasswordProtectedError(
                    f"Word document {file_upload.filename} appears to be password-protected",
                    file_upload.id
                )
            elif "corrupt" in error_str or "invalid" in error_str:
                raise WordCorruptedError(
                    f"Word document {file_upload.filename} appears to be corrupted",
                    file_upload.id
                )
            elif "macro" in error_str or "vba" in error_str:
                raise WordMacroContentError(
                    f"Word document {file_upload.filename} contains macros or VBA content that cannot be processed",
                    file_upload.id
                )
            else:
                raise WordProcessingError(
                    f"python-docx extraction failed: {str(e)}",
                    file_upload.id
                )
    
    def _get_heading_level(self, style_name: str) -> int:
        """Get heading level from style name (Heading 1 -> 1, etc.)."""
        try:
            if 'Heading' in style_name:
                # Extract number from style name like "Heading 1", "Heading 2", etc.
                parts = style_name.split()
                if len(parts) > 1 and parts[1].isdigit():
                    level = int(parts[1])
                    return min(level, 6)  # Max 6 levels for markdown compatibility
                else:
                    return 1  # Default to level 1 if can't parse
            else:
                return 1
        except (ValueError, IndexError):
            return 1
    
    async def create_format_metadata(
        self, 
        file_upload: FileUpload, 
        extracted_content: str
    ) -> Dict[str, Any]:
        """Create Word document-specific metadata."""
        try:
            # Get document for metadata extraction
            word_bytes = await self._get_word_bytes(file_upload)
            doc = Document(io.BytesIO(word_bytes))
            
            # Count structural elements
            paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
            table_count = len(doc.tables)
            
            # Get document properties
            props = doc.core_properties
            
            metadata = {
                "format": "Word Document (.docx)",
                "extraction_library": "python-docx",
                "extracted_text_length": len(extracted_content),
                "paragraph_count": paragraph_count,
                "table_count": table_count,
                "structure_elements": paragraph_count + table_count,
                "document_metadata": {
                    "title": props.title or "",
                    "author": props.author or "",
                    "subject": props.subject or "",
                    "comments": props.comments or "",
                    "keywords": props.keywords or "",
                    "category": props.category or "",
                    "created": props.created.isoformat() if props.created else "",
                    "modified": props.modified.isoformat() if props.modified else "",
                    "last_modified_by": props.last_modified_by or "",
                    "revision": props.revision or "",
                    "version": props.version or ""
                }
            }
            
            # Count headings by level
            heading_counts = {}
            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith('Heading'):
                    level = self._get_heading_level(paragraph.style.name)
                    heading_counts[f"heading_level_{level}"] = heading_counts.get(f"heading_level_{level}", 0) + 1
            
            if heading_counts:
                metadata["heading_structure"] = heading_counts
            
            return metadata
            
        except Exception as e:
            self.logger.warning(f"Error creating Word metadata: {e}")
            return {
                "format": "Word Document",
                "metadata_error": str(e)
            }
